import json
from typing import List, Dict
from docx import Document
from docx.shared import Pt
import os
from fpdf import FPDF

def export_hierarchy_to_txt(hierarchy: Dict, file_handle, indent_level=0):
    """
    Recursively writes the hierarchical structure to the text file.

    Args:
        hierarchy (Dict): Hierarchical dictionary.
        file_handle: Open file handle to write to.
        indent_level (int): Current indentation level.
    """
    indent = '    ' * indent_level
    for key, value in hierarchy.items():
        file_handle.write(f"{indent}{key}/\n")
        if isinstance(value, dict):
            export_hierarchy_to_txt(value, file_handle, indent_level + 1)

def export_to_txt(tree_lines: List[str], skipped_files: List[str], skipped_folders: List[str], output_path: str, hierarchy: Dict = None) -> None:
    """
    Exports the directory tree and skipped items to a text file.
    Optionally includes hierarchical relationships.

    Args:
        tree_lines (List[str]): Lines representing the directory tree.
        skipped_files (List[str]): List of skipped files.
        skipped_folders (List[str]): List of skipped folders.
        output_path (str): Path to the output text file.
        hierarchy (Dict, optional): Hierarchical dictionary to include.
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        if hierarchy:
            f.write("Project Hierarchy:\n")
            f.write("==================\n\n")
            export_hierarchy_to_txt(hierarchy, f)
            f.write("\n")
        else:
            for line in tree_lines:
                f.write(line + '\n')

        # Add summary of skipped items
        if skipped_files or skipped_folders:
            f.write("\nSkipped Items:\n")
            f.write("=" * 20 + "\n")
            
            if skipped_files:
                f.write("\nSkipped Files:\n")
                for file in skipped_files:
                    f.write(f"  {file}\n")
            
            if skipped_folders:
                f.write("\nSkipped Folders:\n")
                for folder in skipped_folders:
                    f.write(f"  {folder}\n")

def export_to_json(tree_hierarchy: Dict, output_path: str) -> None:
    """
    Exports the directory hierarchy to a JSON file.

    Args:
        tree_hierarchy (Dict): Hierarchical dictionary of the directory structure.
        output_path (str): Path to the output JSON file.
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(tree_hierarchy, f, indent=4)

def add_hierarchy_to_docx(doc, hierarchy: Dict, level=0):
    """
    Recursively adds hierarchical data to the Word document.

    Args:
        doc: `python-docx` Document object.
        hierarchy (Dict): Hierarchical dictionary.
        level (int): Current hierarchy level for styling.
    """
    for key, value in hierarchy.items():
        paragraph = doc.add_paragraph(key + '/', style=f'Heading {min(level + 1, 9)}')
        if isinstance(value, dict):
            add_hierarchy_to_docx(doc, value, level + 1)

def export_to_docx(tree_lines: List[str], skipped_files: List[str], skipped_folders: List[str], output_path: str, hierarchy: Dict = None) -> None:
    """
    Exports the directory tree and skipped items to a Word document.
    Optionally includes hierarchical relationships.

    Args:
        tree_lines (List[str]): Lines representing the directory tree.
        skipped_files (List[str]): List of skipped files.
        skipped_folders (List[str]): List of skipped folders.
        output_path (str): Path to the output Word document.
        hierarchy (Dict, optional): Hierarchical dictionary to include.
    """
    doc = Document()
    doc.add_heading('Directory Structure', 0)

    if hierarchy:
        doc.add_heading('Project Hierarchy', level=1)
        add_hierarchy_to_docx(doc, hierarchy)
    else:
        for line in tree_lines:
            if line.endswith('/'):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line, style='List Bullet 2')

    if skipped_files or skipped_folders:
        doc.add_heading('Skipped Items', level=1)
        if skipped_files:
            doc.add_heading('Skipped Files:', level=2)
            for file in skipped_files:
                doc.add_paragraph(file, style='List Bullet 2')
        if skipped_folders:
            doc.add_heading('Skipped Folders:', level=2)
            for folder in skipped_folders:
                doc.add_paragraph(folder, style='List Bullet 2')

    doc.save(output_path)

class PDF(FPDF):
    def header(self):
        # Custom header if needed
        pass

    def footer(self):
        # Custom footer if needed
        pass

def add_hierarchy_to_pdf(pdf: FPDF, hierarchy: Dict, level=0):
    """
    Recursively adds hierarchical data to the PDF.

    Args:
        pdf (FPDF): FPDF object.
        hierarchy (Dict): Hierarchical dictionary.
        level (int): Current hierarchy level for indentation.
    """
    indent = ' ' * (4 * level)
    for key, value in hierarchy.items():
        pdf.set_font("Arial", 'B', 12 - level)
        pdf.multi_cell(0, 10, f"{indent}{key}/")
        if isinstance(value, dict):
            add_hierarchy_to_pdf(pdf, value, level + 1)

def export_to_pdf_direct(hierarchy: Dict, output_path: str) -> None:
    """
    Exports the hierarchical data directly to a PDF.

    Args:
        hierarchy (Dict): Hierarchical dictionary.
        output_path (str): Path to the output PDF file.
    """
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Project Hierarchy", ln=True, align='C')
    pdf.ln(10)

    add_hierarchy_to_pdf(pdf, hierarchy)

    pdf.output(output_path)
